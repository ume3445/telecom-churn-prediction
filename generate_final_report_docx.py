from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def set_global_format(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if level > 1 else 14)
    p.paragraph_format.line_spacing = 1.5


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_para(doc, f"[Missing figure file: {image_path.name}]")
        return
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_run = pic.add_run()
    pic_run.add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.line_spacing = 1.5


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr_cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")


def main() -> None:
    out_dir = Path("outputs")
    out_file = out_dir / "CSC398_Final_Report.docx"

    model_df = pd.read_csv(out_dir / "model_comparison_results.csv")
    thr_enet = pd.read_csv(out_dir / "threshold_analysis_best_model.csv")
    thr_ridge = pd.read_csv(out_dir / "threshold_analysis_interpretable_model.csv")
    cal_df = pd.read_csv(out_dir / "calibration_results.csv")
    shap_df = pd.read_csv(out_dir / "shap/top10_mean_abs_shap.csv")
    uplift_seg = pd.read_csv(out_dir / "uplift/segment_counts.csv")
    cox_df = pd.read_csv(out_dir / "survival/cox_top10_hazard_ratios.csv")
    roi_df = pd.read_csv(out_dir / "roi/roi_simulation_grid.csv")

    abstract = (
        "Telecommunication providers face persistent revenue erosion from customer churn, yet many predictive systems remain disconnected from intervention design and economic decision support. "
        "We develop an interpretable churn intelligence framework using the Telco Customer Churn dataset with 7,043 customers and 33 variables. "
        "Our baseline modeling stage compares class-balanced logistic regression, Ridge, Lasso, and Elastic Net under a stratified 80/20 split and five-fold cross-validation. "
        "The best predictive model is calibrated Elastic Net, achieving a test ROC AUC of 0.8484 and a calibrated Brier score of 0.1361. "
        "We extend this baseline with four novelty modules: SHAP-based instance-level explainability, uplift modeling via a T-learner, survival analysis for time-to-churn estimation, and revenue-at-risk plus ROI simulation. "
        "SHAP shows Tenure_Months as the strongest global contributor with mean absolute SHAP value 0.9162. "
        "The uplift segmentation identifies 474 Persuadables and 231 Sure Things, enabling intervention triage beyond risk ranking. "
        "Cox survival modeling yields a test concordance index of 0.9028, supporting actionable churn timing estimates. "
        "At a 0.4 threshold, annualized revenue at risk is $242,525.92, and threshold-specific campaign simulations show positive net ROI across all tested policies. "
        "The resulting framework integrates interpretability, causally aware targeting, temporal risk, and financial prioritization for deployable retention strategy in modern telecom operations at scale."
    )
    if len(abstract.split()) != 200:
        raise ValueError(f"Abstract word count is {len(abstract.split())}, expected 200.")

    doc = Document()
    set_global_format(doc)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Interpretable Churn Prediction for Telecom Decision Support: Regularized Logistic Models, Uplift Modeling, and Revenue-Aware Retention Optimization"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    for line in [
        "",
        "Moeez Islam Malik",
        "Muhammad Umer Hammad",
        "Muhammad Abdullah Saleem",
        "",
        "CSC 398 - Independent Study",
        "Instructor: Professor Mehmet",
        "Spring 2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_para(doc, abstract)

    add_heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "Customer churn is a structural challenge in the telecommunications sector because recurring revenue models amplify the financial impact of each lost subscriber. "
        "As competitive switching costs decline and service parity increases, providers require predictive systems that identify not only likely churners but also operationally meaningful intervention opportunities. "
        "Many production workflows still rely on static risk ranking, which captures correlation with churn yet often omits treatment responsiveness, temporal urgency, and value-weighted targeting."
    )
    add_para(
        doc,
        "This study addresses that gap by extending a calibrated regularized logistic baseline into a decision-support architecture that links prediction to action. "
        "We first establish a strong interpretable benchmark with Elastic Net and probability calibration, then introduce four contributions. "
        "Contribution 1 is SHAP-based explanation at both global and individual levels, allowing analyst and manager interpretation beyond coefficient summaries. "
        "Contribution 2 is uplift modeling that estimates incremental intervention impact and separates persuadable customers from non-actionable cohorts. "
        "Contribution 3 is survival analysis that reframes churn as a time-to-event process and supports intervention scheduling. "
        "Contribution 4 is revenue-at-risk and return-on-investment simulation to align model outputs with campaign economics."
    )
    add_para(
        doc,
        "The remainder of the paper proceeds from related literature to dataset and baseline methods, then presents each novelty module with empirical findings and managerial implications. "
        "We conclude with an integrated deployment framework, limitations, and future work toward experimentally validated causal retention optimization."
    )

    add_heading(doc, "1.1 Novelty Positioning", 2)
    add_para(
        doc,
        "The project's novelty is primarily integrative rather than algorithmic. Prior literature has independently studied churn prediction, SHAP-based explainability, uplift modeling, survival analysis, and campaign economics. "
        "However, most studies do not operationalize these components in a unified telecom decision pipeline anchored to a calibrated base model. "
        "Our contribution is the end-to-end coupling of discrimination, interpretability, intervention response estimation, time-to-churn modeling, and revenue-aware policy simulation within one reproducible workflow."
    )
    add_para(
        doc,
        "Accordingly, we position this work as methodological integration and operationalization novelty. "
        "This framing is academically precise: we do not claim a new learning algorithm, but we do contribute a practically deployable architecture that addresses critical gaps left by single-module churn studies."
    )

    add_heading(doc, "2. Related Work", 1)
    add_para(
        doc,
        "Prior churn research demonstrates that logistic regression remains competitive when feature engineering and calibration are handled rigorously, while retaining interpretability that is often absent in higher-capacity black-box systems. "
        "Classical marketing analytics work established customer lifetime risk and retention propensity modeling as a core mechanism for subscription businesses, and subsequent telecom studies adapted these methods to service-level and billing behavior features."
    )
    add_para(
        doc,
        "Explainable AI literature provides methods that localize model reasoning to individual predictions. "
        "SHAP, grounded in Shapley values from cooperative game theory, offers additive attributions that satisfy local accuracy and consistency properties, making it suitable for audit-sensitive churn operations where case-level rationale is required."
    )
    add_para(
        doc,
        "Uplift modeling extends beyond risk prediction by estimating differential treatment response, which has been widely used in direct marketing and policy targeting. "
        "This distinction is critical in retention settings because high-risk customers are not necessarily the customers most likely to respond to intervention."
    )
    add_para(
        doc,
        "Survival analysis has also matured in customer analytics, providing interpretable hazard and survival-time estimates for attrition processes. "
        "Kaplan-Meier and Cox models offer transparent temporal dynamics that can be operationalized in contact scheduling. "
        "Despite this progress, prior work often examines these strands in isolation rather than integrating discrimination, explanation, uplift, timing, and economic optimization in one end-to-end telecom pipeline, which is the central contribution of this project."
    )

    add_heading(doc, "3. Dataset and Data Preparation", 1)
    add_para(
        doc,
        "We use the Telco Customer Churn benchmark dataset containing 7,043 observations and 33 columns spanning demographics, subscribed services, contract and billing behavior, and churn outcomes. "
        "The target is binary churn status, with approximately 26.5% churners and 73.5% non-churners, creating moderate class imbalance that motivates balanced learning and threshold-aware evaluation."
    )
    add_para(
        doc,
        "Data preparation standardized schema naming, coerced Total_Charges into numeric form, and handled tenure-linked edge cases in cumulative billing values. "
        "To prevent leakage, we excluded identifiers and post-outcome signals, including CustomerID, Churn_Label, Churn_Score, and Churn_Reason. "
        "We used a stratified 80/20 split with random_state=42 and embedded imputation, scaling, and one-hot encoding inside a scikit-learn pipeline so transformations were always learned only from training folds."
    )

    add_heading(doc, "4. Baseline Modeling and Evaluation", 1)
    add_para(
        doc,
        "The baseline stage compared class-balanced logistic regression variants: unregularized baseline, Ridge (L2), Lasso (L1), and Elastic Net. "
        "Each model used consistent preprocessing and was evaluated with five-fold cross-validation and held-out test metrics, including ROC AUC, PR AUC, precision, recall, F1-score, and Brier score."
    )
    add_para(
        doc,
        "Elastic Net achieved the highest test ROC AUC at 0.8484, while Ridge and Lasso were very close at 0.8477 and 0.8476. "
        "Baseline logistic reached 0.8366, confirming the value of regularization under sparse one-hot feature expansion. "
        "Calibration materially improved reliability: isotonic-calibrated Elastic Net reduced Brier score from 0.1641 to 0.1361, supporting downstream probability-based intervention budgeting."
    )
    add_para(
        doc,
        "Threshold analysis reinforced that 0.5 is not universally optimal in retention operations. "
        "For Elastic Net, reducing threshold from 0.5 to 0.4 increased recall from 0.7807 to 0.8690 at expected precision cost, reflecting a tractable policy trade-off where missing churners is expensive. "
        "These results support selecting calibrated Elastic Net as the production scoring core before extending to novelty modules."
    )
    add_figure(doc, out_dir / "roc_curves.png", "Figure 1. ROC curves for baseline and regularized logistic models.")
    add_figure(doc, out_dir / "pr_curves.png", "Figure 2. Precision-Recall curves for baseline and regularized logistic models.")
    add_figure(doc, out_dir / "calibration_curves.png", "Figure 3. Calibration curves for raw and calibrated top models.")

    add_heading(doc, "5. SHAP-Based Explainability", 1)
    add_para(
        doc,
        "Coefficient inspection is useful for global directionality, yet it cannot fully explain non-uniform local feature contributions after preprocessing and regularization interactions. "
        "We therefore implemented SHAP LinearExplainer on the calibrated Elastic Net base estimator using transformed training data, then computed test-set SHAP values for global and case-level interpretation."
    )
    add_para(
        doc,
        "Global SHAP analysis identified Tenure_Months as the strongest driver with mean absolute SHAP value 0.9162. "
        "Dependents_Yes, Contract_Two_year, Internet_Service_Fiber_optic, and Contract_One_year followed, confirming that relationship maturity and contract structure dominate churn propensity. "
        "Waterfall decompositions for the three highest-risk customers showed heterogeneous combinations of short tenure, contract type, service mix, and payment behavior, demonstrating why individual intervention narratives are more nuanced than a single coefficient ranking."
    )
    add_para(
        doc,
        "SHAP rankings aligned with coefficient-based findings while adding case-specific magnitude and sign context, which improves analyst trust and frontline actionability. "
        "In operational settings, this supports transparent customer-contact rationale and improves compliance readiness for model-assisted decision systems."
    )
    add_figure(doc, out_dir / "shap/shap_summary_beeswarm_top15.png", "Figure 4. SHAP beeswarm summary plot for top 15 features by mean absolute attribution.")
    add_figure(doc, out_dir / "shap/waterfall_top_risk_1.png", "Figure 5. SHAP waterfall plot for the highest-risk test customer.")

    add_heading(doc, "6. Uplift Modeling", 1)
    add_para(
        doc,
        "Churn prediction estimates who is likely to leave, whereas uplift modeling estimates who is likely to change outcome under intervention. "
        "This distinction matters because retention resources are scarce and some high-risk customers are either inherently retainable or effectively non-responsive."
    )
    add_para(
        doc,
        "Because the benchmark dataset does not include treatment logs, we implemented a transparent simulation for method development. "
        "We assigned treatment to approximately half of training customers, simulated average churn-risk reduction near 25% with controlled noise, and trained a T-learner using separate treated and control logistic models. "
        "We then scored the untouched test set with uplift estimates defined as P(churn|treated) minus P(churn|control), where more negative values imply greater expected treatment benefit."
    )
    add_para(
        doc,
        "The resulting segmentation yielded 474 Persuadables, 473 Lost Causes, 231 Sure Things, and 231 Do Not Disturb customers. "
        "This distribution shows that risk ranking alone would over-target non-incremental cohorts. "
        "The Qini-style cumulative gains curve further indicated stronger incremental benefit when targeting by modeled uplift than random contact policy. "
        "Managerially, these results support prioritizing Persuadables and de-prioritizing Sure Things and Do Not Disturb segments when budget-constrained."
    )
    add_figure(doc, out_dir / "uplift/segmentation_2x2.png", "Figure 6. Uplift segmentation map: churn risk by estimated uplift potential.")
    add_figure(doc, out_dir / "uplift/qini_curve.png", "Figure 7. Qini-style cumulative gains curve for uplift-ranked targeting.")

    add_heading(doc, "7. Survival Analysis", 1)
    add_para(
        doc,
        "Binary churn classifiers suppress timing information that is central to intervention scheduling. "
        "We reframed churn as time-to-event using tenure in months as duration and churn outcome as event indicator, then fitted Kaplan-Meier, Cox proportional hazards, and Weibull AFT models on the training partition."
    )
    add_para(
        doc,
        "Kaplan-Meier curves stratified by contract type showed clear separation, with month-to-month customers exhibiting faster survival decline than one-year and two-year cohorts. "
        "The Cox model achieved a test concordance index of 0.9028, indicating strong ranking quality for ordering earlier versus later churn events. "
        "For the five highest-risk customers from calibrated Elastic Net, predicted median survival ranged from 3.24 to 11.05 months, enabling proactive timing of retention outreach windows."
    )
    add_para(
        doc,
        "Temporal risk estimates complement probability scores by answering when intervention should occur, not only whether intervention should occur. "
        "This improves campaign cadence design and can reduce both premature and delayed outreach."
    )
    add_figure(doc, out_dir / "survival/km_by_contract.png", "Figure 8. Kaplan-Meier survival curves stratified by contract type.")
    add_figure(doc, out_dir / "survival/cox_hazard_forest_top10.png", "Figure 9. Cox proportional hazards forest plot for top ten covariates.")

    add_heading(doc, "8. Revenue-at-Risk and ROI Simulation", 1)
    add_para(
        doc,
        "To align model outputs with business outcomes, we computed customer-level annualized revenue-at-risk as calibrated churn probability multiplied by monthly charges and twelve months. "
        "This transformed risk scores into directly interpretable exposure estimates."
    )
    add_para(
        doc,
        "At threshold 0.4, total revenue at risk reached $242,525.92 in the held-out test set. "
        "We then simulated campaign economics across thresholds from 0.2 to 0.6 with default assumptions of $50 offer cost, 30% retention success among caught churners, $65 monthly revenue, and twelve retained months. "
        "All policies produced positive net ROI, with net values from $25,336 to $41,932 and increasing ROI percentage at higher thresholds because contact volume declines faster than expected retained value."
    )
    add_para(
        doc,
        "These results suggest policy selection should trade absolute retained revenue against contact intensity and operational capacity. "
        "For organizations prioritizing balanced scale and return, thresholds around 0.3 to 0.4 are practical starting points."
    )
    add_figure(doc, out_dir / "roi/revenue_at_risk_scatter.png", "Figure 10. Revenue-at-risk map combining calibrated churn probability and monthly charges.")
    add_figure(doc, out_dir / "roi/net_roi_by_threshold.png", "Figure 11. Net ROI under threshold-specific retention policies.")

    add_heading(doc, "9. Integrated Framework and Business Recommendations", 1)
    add_para(
        doc,
        "The combined architecture supports an end-to-end retention workflow: score risk with calibrated Elastic Net, explain decisions with SHAP, filter contact candidates by uplift potential, prioritize by revenue-at-risk, and schedule intervention based on survival timing. "
        "This sequence links statistical discrimination to intervention incrementality, financial prioritization, and campaign timing."
    )
    add_para(
        doc,
        "For deployment, we recommend batch scoring integrated with CRM automation, decision thresholds parameterized by budget, and policy-level monitoring of precision, uplift realization, and realized ROI. "
        "Human-in-the-loop review should remain mandatory for high-impact segments, both to audit atypical explanations and to align interventions with regulatory and customer-experience constraints."
    )

    add_heading(doc, "10. Limitations", 1)
    add_para(
        doc,
        "The study uses a benchmark dataset that may not capture provider-specific market dynamics, channel effects, or temporal drift. "
        "The uplift component relies on simulated treatment, which is useful for methodological prototyping but does not replace causal identification from real randomized exposure."
    )
    add_para(
        doc,
        "Model families are primarily linear in log-odds or semi-parametric hazard form, so nonlinear interactions may remain underrepresented. "
        "We also do not perform temporal external validation on a future cohort, which limits inference about post-deployment stability. "
        "Finally, causal conclusions about intervention effectiveness remain provisional until tested with controlled experiments."
    )

    add_heading(doc, "11. Conclusion", 1)
    add_para(
        doc,
        "This paper contributes a unified interpretable churn decision framework that combines calibrated regularized logistic prediction, SHAP local-global explanation, uplift-based actionability, survival-based timing, and value-aware economic optimization. "
        "The empirical core achieved ROC AUC 0.8484 with calibrated Brier score 0.1361, produced uplift segmentation with 474 Persuadables, reached Cox C-index 0.9028, and identified $242,525.92 in revenue-at-risk at threshold 0.4."
    )
    add_para(
        doc,
        "Future work should replace simulated treatment with randomized A/B retention experiments, add drift monitoring and recalibration services, incorporate real intervention cost heterogeneity, and benchmark against non-linear and deep learning alternatives under identical business constraints."
    )

    add_heading(doc, "References", 1)
    references = [
        "Burez, J., & Van den Poel, D. (2009). Handling class imbalance in customer churn prediction. Expert Systems with Applications, 36(3), 4626-4636. https://doi.org/10.1016/j.eswa.2008.05.027",
        "Buckinx, W., & Van den Poel, D. (2005). Customer base analysis: Partial defection of behaviorally loyal clients in a non-contractual FMCG retail setting. European Journal of Operational Research, 164(1), 252-268. https://doi.org/10.1016/j.ejor.2003.12.010",
        "Coussement, K., & Van den Poel, D. (2008). Churn prediction in subscription services: An application of support vector machines while comparing two parameter-selection techniques. Expert Systems with Applications, 34(1), 313-327. https://doi.org/10.1016/j.eswa.2006.09.038",
        "De Caigny, A., Coussement, K., & De Bock, K. W. (2018). A new hybrid classification algorithm for customer churn prediction based on logistic regression and decision trees. European Journal of Operational Research, 269(2), 760-772. https://doi.org/10.1016/j.ejor.2018.02.009",
        "Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). Applied logistic regression (3rd ed.). Wiley.",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. In Advances in Neural Information Processing Systems (pp. 4765-4774).",
        "Radcliffe, N. J., & Surry, P. D. (2011). Real-world uplift modelling with significance-based uplift trees. White paper, Stochastic Solutions.",
        "Rzepakowski, P., & Jaroszewicz, S. (2012). Decision trees for uplift modeling with single and multiple treatments. Knowledge and Information Systems, 32, 303-327. https://doi.org/10.1007/s10115-011-0436-0",
        "Vapnik, V. (1998). Statistical learning theory. Wiley.",
        "Vink, M. A., & Van den Hout, A. (2019). Calibration of risk prediction models in survival analysis. Statistics in Medicine, 38(11), 1969-1982. https://doi.org/10.1002/sim.8083",
        "Verbraken, T., Verbeke, W., & Baesens, B. (2014). A novel profit maximizing metric for measuring classification performance of customer churn prediction models. IEEE Transactions on Knowledge and Data Engineering, 25(5), 961-973. https://doi.org/10.1109/TKDE.2012.50",
        "Witten, I. H., Frank, E., Hall, M. A., & Pal, C. J. (2016). Data mining: Practical machine learning tools and techniques (4th ed.). Morgan Kaufmann.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()
    add_heading(doc, "Appendices", 1)
    add_table_from_df(doc, model_df, "Appendix A. Full Model Comparison Table")
    add_table_from_df(doc, thr_enet, "Appendix B1. Threshold Analysis (Elastic Net)")
    add_table_from_df(doc, thr_ridge, "Appendix B2. Threshold Analysis (Ridge)")
    add_table_from_df(doc, cal_df, "Appendix C. Calibration Results")
    add_table_from_df(doc, shap_df, "Appendix D. Top 10 SHAP Features")
    add_table_from_df(doc, uplift_seg, "Appendix E. Uplift Segment Counts")
    add_table_from_df(doc, cox_df, "Appendix F. Cox Hazard Ratios")
    add_table_from_df(doc, roi_df, "Appendix G. ROI Simulation Grid")

    doc.save(out_file)
    print(f"Saved report to: {out_file.resolve()}")


if __name__ == "__main__":
    main()
