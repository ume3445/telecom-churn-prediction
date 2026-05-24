from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INPUT = Path("outputs/CSC398_Final_Report.docx")
OUTPUT = Path("outputs/CSC398_Final_Report_FINAL.docx")


H1_TEXTS = {
    "Abstract",
    "1. Introduction",
    "2. Related Work",
    "3. Dataset and Data Preparation",
    "4. Baseline Modeling and Evaluation",
    "5. SHAP-Based Explainability",
    "6. Uplift Modeling",
    "7. Survival Analysis",
    "8. Revenue-at-Risk and ROI Simulation",
    "9. Integrated Framework and Business Recommendations",
    "10. Limitations",
    "11. Conclusion",
    "References",
    "Appendices",
}

H2_TEXTS = {
    "Appendix A. Full Model Comparison Table",
    "Appendix B1. Threshold Analysis (Elastic Net)",
    "Appendix B2. Threshold Analysis (Ridge)",
    "Appendix C. Calibration Results",
    "Appendix D. Top 10 SHAP Features",
    "Appendix E. Uplift Segment Counts",
    "Appendix F. Cox Hazard Ratios",
    "Appendix G. ROI Simulation Grid",
}


def set_paragraph_border(paragraph, **kwargs):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    for edge in ("top", "left", "bottom", "right", "between"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = pBdr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            pBdr.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def add_page_number(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11 pt
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    run.append(t)
    fld.append(run)
    paragraph._p.append(fld)


def ensure_two_sections(doc: Document):
    if len(doc.sections) >= 2:
        return
    # Insert section break at title-page last paragraph (index 9 in current structure).
    body = doc._element.body
    sect_pr = body.sectPr
    split_para = doc.paragraphs[9]._p
    pPr = split_para.get_or_add_pPr()
    p_sect = deepcopy(sect_pr)
    pPr.append(p_sect)
    # Mark following section as next page.
    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sect_pr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")


def set_run_font(run, size_pt=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def main():
    doc = Document(str(INPUT))
    original_text = "\n".join(p.text for p in doc.paragraphs)
    original_images = len(doc.inline_shapes)
    original_tables = len(doc.tables)

    ensure_two_sections(doc)

    # Global margins
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Heading style definitions
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(13)
    h2.font.bold = True

    # Apply paragraph-level styles and fonts
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt in H1_TEXTS:
            p.style = doc.styles["Heading 1"]
            for r in p.runs:
                set_run_font(r, size_pt=14, bold=True, italic=False)
            continue
        if txt in H2_TEXTS:
            p.style = doc.styles["Heading 2"]
            for r in p.runs:
                set_run_font(r, size_pt=13, bold=True, italic=False)
            continue
        if txt.startswith("Figure"):
            for r in p.runs:
                set_run_font(r, size_pt=11, bold=False, italic=True)
            continue
        if 0 <= i <= 9:
            continue
        for r in p.runs:
            set_run_font(r, size_pt=12, bold=False, italic=False)

    # Header/footer only from page 2 onwards
    if len(doc.sections) >= 2:
        sec1 = doc.sections[0]
        sec2 = doc.sections[1]
        sec1.different_first_page_header_footer = True
        sec2.different_first_page_header_footer = False
        sec2.header.is_linked_to_previous = False
        sec2.footer.is_linked_to_previous = False

        for p in sec1.header.paragraphs:
            p.text = ""
        for p in sec1.footer.paragraphs:
            p.text = ""

        hp = sec2.header.paragraphs[0] if sec2.header.paragraphs else sec2.header.add_paragraph()
        hp.text = "Churn Prediction Decision Support | CSC 398"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in hp.runs:
            set_run_font(r, size_pt=10, color="808080")

        fp = sec2.footer.paragraphs[0] if sec2.footer.paragraphs else sec2.footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp)
    else:
        # Fallback: different first page on single section.
        sec = doc.sections[0]
        sec.different_first_page_header_footer = True
        hp = sec.header.paragraphs[0]
        hp.text = "Churn Prediction Decision Support | CSC 398"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in hp.runs:
            set_run_font(r, size_pt=10, color="808080")
        fp = sec.footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp)

    # Title page redesign (paragraph indices based on existing document)
    title_p = doc.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(24)
    title_p.paragraph_format.line_spacing = 1.3
    set_paragraph_border(title_p, top={"val": "single", "sz": 24, "space": 1, "color": "1F3864"})
    for r in title_p.runs:
        set_run_font(r, size_pt=20, bold=True, color="1F3864")

    divider_p = doc.paragraphs[1]
    divider_p.text = ""
    divider_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_p.paragraph_format.left_indent = Inches(2.25)
    divider_p.paragraph_format.right_indent = Inches(2.25)
    divider_p.paragraph_format.space_after = Pt(24)
    set_paragraph_border(divider_p, bottom={"val": "single", "sz": 24, "space": 1, "color": "C9A84C"})

    for idx in [2, 3, 4]:
        p = doc.paragraphs[idx]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            set_run_font(r, size_pt=13, bold=False, color="404040")

    course_p = doc.paragraphs[6]
    course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_p.paragraph_format.space_before = Pt(36)
    for r in course_p.runs:
        set_run_font(r, size_pt=11, bold=True, color="1F3864")

    for idx in [7, 8]:
        p = doc.paragraphs[idx]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run_font(r, size_pt=11, bold=False, color="606060")

    bottom_bar_p = doc.paragraphs[9]
    bottom_bar_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_bar_p.paragraph_format.space_before = Pt(80)
    set_paragraph_border(bottom_bar_p, top={"val": "single", "sz": 24, "space": 1, "color": "1F3864"})

    # Save and verify no content/image/table loss.
    doc.save(str(OUTPUT))

    out_doc = Document(str(OUTPUT))
    new_text = "\n".join(p.text for p in out_doc.paragraphs)
    assert original_text == new_text, "Text content changed."
    assert len(out_doc.inline_shapes) == original_images, "Image count changed."
    assert len(out_doc.tables) == original_tables, "Table count changed."

    print(f"Saved: {OUTPUT}")
    print(f"Images: {len(out_doc.inline_shapes)} | Tables: {len(out_doc.tables)} | Sections: {len(out_doc.sections)}")


if __name__ == "__main__":
    main()
